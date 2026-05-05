from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT_PATH = Path("/Users/python/Documents/Public AI Trial /Outputs/DEWR_Public_AI_B.docx")

DEWR_GREEN = "719F4C"
DEWR_DARK_GREEN = "5D7A38"
DEWR_DARK_GREY = "404246"
DEWR_GREY = "A4A7A9"
DEWR_LIGHT_GREY = "D7D8D8"
PANEL_GREY = "F7F8FA"
TRACK_GREY = "E3E5E6"
SOFT_GREEN = "EFF4E8"
WHITE = "FFFFFF"


def rgb(hex_color):
    h = hex_color.replace("#", "")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_width(table, pct=5000):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), str(pct))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True


def clear_cell(cell):
    for p in cell.paragraphs:
        p.clear()


def add_run(paragraph, text, bold=False, italic=False, size=None, color=DEWR_DARK_GREY):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.font.name = "Aptos"
    return run


def add_paragraph(doc, text="", style=None, bold=False, italic=False, size=None, color=DEWR_DARK_GREY,
                  align=None, space_after=6):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        add_run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def add_cell_text(cell, text, bold=False, italic=False, size=9, color=DEWR_DARK_GREY,
                  align=None, space_after=0):
    clear_cell(cell)
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    add_run(p, text, bold=bold, italic=italic, size=size, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    return p


def add_rule(paragraph, color=DEWR_GREEN, size=8):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = rgb(DEWR_DARK_GREY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, before, after in [
        ("Title", 26, 0, 6),
        ("Heading 1", 18, 16, 8),
        ("Heading 2", 14, 14, 6),
        ("Heading 3", 11.5, 10, 4),
    ]:
        style = styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(DEWR_DARK_GREY)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def setup_doc():
    doc = Document()
    set_styles(doc)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    add_run(hp, "Evaluation of the Public Generative AI Trial", size=8, color=DEWR_DARK_GREY)
    add_rule(hp, color=DEWR_DARK_GREY, size=12)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(fp, "Page ", size=8, color=DEWR_DARK_GREY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    fp._p.append(fld_begin)
    fp._p.append(instr)
    fp._p.append(fld_end)
    return doc


def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def heading(doc, text, level=1, rule=False):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    if rule:
        add_rule(p, color=DEWR_GREEN, size=6)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_run(p, text, size=10)


def callout(doc, text, fill=DEWR_DARK_GREEN):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=160, bottom=160, start=220, end=220)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, text, bold=True, size=10.5, color=WHITE)
    add_paragraph(doc, "", space_after=4)
    return table


def key_finding(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, SOFT_GREEN)
    set_cell_border(cell, left={"val": "single", "sz": 18, "color": DEWR_GREEN})
    set_cell_margins(cell, top=110, bottom=110, start=220, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_run(p, text, bold=True, size=9.4)
    add_paragraph(doc, "", space_after=2)


def visual_title(doc, text):
    add_paragraph(doc, text, bold=True, size=10.5, space_after=3)


def source_note(doc, text):
    add_paragraph(doc, text, italic=True, size=8.2, color=DEWR_GREY, space_after=6)


def stat_cards(doc, title, cards, note=None):
    visual_title(doc, title)
    table = doc.add_table(rows=1, cols=len(cards))
    set_table_width(table)
    for i, (value, label, color) in enumerate(cards):
        cell = table.cell(0, i)
        set_cell_shading(cell, color)
        set_cell_margins(cell, top=140, bottom=130, start=120, end=120)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        add_run(p, value, bold=True, size=20, color=WHITE)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        add_run(p2, label, size=8.4, color=WHITE)
    if note:
        source_note(doc, note)


def value_panel(doc, title, items, primary_count=1, note=None):
    visual_title(doc, title)
    table = doc.add_table(rows=1, cols=len(items))
    set_table_width(table)
    for i, (value, label) in enumerate(items):
        cell = table.cell(0, i)
        set_cell_shading(cell, PANEL_GREY)
        set_cell_border(cell, right={"val": "single", "sz": 4, "color": DEWR_LIGHT_GREY})
        set_cell_margins(cell, top=170, bottom=150, start=120, end=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        add_run(p, value, bold=True, size=20, color=DEWR_GREEN if i < primary_count else DEWR_DARK_GREY)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        add_run(p2, label, size=7.7)
    if note:
        source_note(doc, note)


def matrix_panel(doc, title, columns, rows, highlight_col=None, note=None):
    visual_title(doc, title)
    table = doc.add_table(rows=1, cols=len(columns))
    set_table_width(table)
    hdr = table.rows[0]
    for i, col in enumerate(columns):
        cell = hdr.cells[i]
        set_cell_shading(cell, PANEL_GREY)
        set_cell_border(cell, bottom={"val": "single", "sz": 4, "color": DEWR_LIGHT_GREY})
        add_cell_text(cell, col.upper(), bold=True, size=7.6, align=WD_ALIGN_PARAGRAPH.CENTER if i else None)
    set_repeat_table_header(hdr)
    for row in rows:
        cells = table.add_row().cells
        label = row[0]
        values = row[1:]
        add_cell_text(cells[0], label, bold=True, size=8.2)
        for i, value in enumerate(values, start=1):
            add_cell_text(
                cells[i],
                str(value),
                bold=True,
                size=10,
                color=DEWR_DARK_GREEN if highlight_col is not None and i == highlight_col else DEWR_DARK_GREY,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
        for cell in cells:
            set_cell_shading(cell, PANEL_GREY)
            set_cell_border(cell, bottom={"val": "single", "sz": 3, "color": DEWR_LIGHT_GREY})
    if note:
        source_note(doc, note)


def bar_panel(doc, title, header, items, primary_count=1, note=None, segments=24):
    visual_title(doc, title)
    table = doc.add_table(rows=1, cols=3)
    set_table_width(table)
    row = table.rows[0]
    add_cell_text(row.cells[0], header.upper(), bold=True, size=7.5)
    set_cell_shading(row.cells[0], PANEL_GREY)
    add_cell_text(row.cells[1], "", size=7.5)
    add_cell_text(row.cells[2], "", size=7.5)
    for cell in row.cells:
        set_cell_shading(cell, PANEL_GREY)
    for i, (label, value) in enumerate(items):
        row = table.add_row()
        cells = row.cells
        add_cell_text(cells[0], label, size=7.8)
        set_cell_shading(cells[0], PANEL_GREY)
        clear_cell(cells[1])
        p = cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        filled = round((value / 100) * segments)
        filled_color = DEWR_DARK_GREEN if i < primary_count else DEWR_DARK_GREY
        add_run(p, "■" * filled, size=8.5, color=filled_color)
        add_run(p, "■" * (segments - filled), size=8.5, color=TRACK_GREY)
        p.paragraph_format.space_after = Pt(0)
        set_cell_shading(cells[1], PANEL_GREY)
        set_cell_margins(cells[1], top=80, bottom=80, start=80, end=80)
        add_cell_text(cells[2], f"{value:g}%", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_shading(cells[2], PANEL_GREY)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})
    if note:
        source_note(doc, note)


def concern_panel(doc):
    visual_title(doc, "Open-text concerns focused on access, data handling and accuracy")
    table = doc.add_table(rows=4, cols=2)
    set_table_width(table)
    headers = ["COMMON PRACTICAL CONCERNS", "BROADER TRUST AND IMPACT CONCERNS"]
    data = [
        ("*** Tool access, integration and cost (17 mentions)", "*-- Environmental impact (5 mentions)"),
        ("**- Data privacy, confidentiality and public-tool data use (13)", "*-- Workforce, capability and civic reliance (4)"),
        ("*-- Accuracy, hallucination and validation (8)", "*-- Other: bias, ethics, governance (2)"),
    ]
    for c, header in enumerate(headers):
        cell = table.cell(0, c)
        set_cell_shading(cell, PANEL_GREY)
        add_cell_text(cell, header, bold=True, size=7.7)
    for r, row_data in enumerate(data, start=1):
        for c, text in enumerate(row_data):
            cell = table.cell(r, c)
            set_cell_shading(cell, PANEL_GREY)
            display = text.replace("*", "●").replace("-", "○")
            add_cell_text(cell, display, size=8)
    source_note(
        doc,
        "Note: More filled dots indicate themes raised more often in open-text responses (n=51; 32 open-text entries). "
        "Counts are mentions, not unique respondents, and should not be summed. Source: DEWR Public Generative AI Trial survey, 2026.",
    )


def judgement_panel(doc):
    visual_title(doc, "Staff judgement was needed at three practical boundaries")
    table = doc.add_table(rows=2, cols=3)
    set_table_width(table)
    headers = [
        "Information classification boundaries",
        "Output validation",
        "APS-specific sensitivity and safeguards",
    ]
    body = [
        "Users described uncertainty about what information was appropriate to enter.",
        "Users pointed to the need to check outputs before relying on them.",
        "Some responses highlighted differences between public tools and Copilot safeguards.",
    ]
    for c in range(3):
        set_cell_shading(table.cell(0, c), PANEL_GREY)
        add_cell_text(table.cell(0, c), headers[c], bold=True, size=9)
        set_cell_shading(table.cell(1, c), PANEL_GREY)
        add_cell_text(table.cell(1, c), body[c], size=8.1)


def add_quote(doc, text):
    p = add_paragraph(doc, f'"{text}"', italic=True, size=9, color=DEWR_GREY, space_after=5)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)


def build_doc():
    doc = setup_doc()

    doc.add_paragraph()
    title = doc.add_paragraph(style="Title")
    add_run(title, "Evaluation of the Public\nGenerative AI Trial", bold=True, size=26)
    subtitle = doc.add_paragraph()
    add_run(subtitle, "Summary of evaluation findings", size=14, color=DEWR_DARK_GREY)
    add_rule(subtitle, color=DEWR_GREEN, size=6)
    add_paragraph(doc, "Department of Employment and Workplace Relations", size=11)
    add_paragraph(doc, "May 2026", size=10, color=DEWR_GREY)
    page_break(doc)

    heading(doc, "Contents", 1)
    toc = [
        ("Executive summary", "3"),
        ("Preface", "3"),
        ("Overarching findings", "4"),
        ("Recommendations", "5"),
        ("Evaluation findings", "6"),
        ("1. Copilot productivity baseline", "6"),
        ("2. Public Gen AI uptake, use and productivity", "8"),
        ("3. Concerns, risks and safeguards", "11"),
        ("Approach and methodology", "15"),
        ("Appendix", "16"),
    ]
    t = doc.add_table(rows=0, cols=2)
    set_table_width(t)
    for item, page in toc:
        cells = t.add_row().cells
        add_cell_text(cells[0], item, bold=item in {"Executive summary", "Evaluation findings", "Approach and methodology", "Appendix"}, size=10)
        add_cell_text(cells[1], page, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
        for cell in cells:
            set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})
    page_break(doc)

    heading(doc, "Executive summary", 1, rule=True)
    heading(doc, "Preface", 2)
    add_paragraph(doc, "DEWR currently provides all employees with access to Generative AI. All employees can access the free Microsoft Copilot Chat and just under 10% can access the paid M365 Copilot with greater integration capability. In March 2026, nearly 80% of all staff used one of these tools, entering over 250,000 prompts, or about 4 prompts per person every workday.")
    add_paragraph(doc, "In January 2026, DEWR ran a trial in which 5% of employees were provided access to Public Generative AI tools, in addition to their existing Copilot access. The trial cohort was selected at random and stratified to ensure representation across Groups, Copilot access types, and APS levels. The tools were the free versions of OpenAI's ChatGPT, Google's Gemini, and Anthropic's Claude, accessed via web browsers. Staff were provided with technical and governance instruction, including on what could be uploaded (unclassified information only). Technical protections were established to reduce the likelihood of classified material being uploaded to the tools.")
    add_paragraph(doc, "At the conclusion of the trial period, all participants were invited to complete a voluntary survey about their experience. A total of 104 staff completed the survey, representing an approximate 52% response rate. Respondents were evenly split by classification, with 52 APS staff (50%) and 52 Executive Level staff (50%), drawn from all five departmental groups.")
    add_paragraph(doc, "The trial was intended to assess:", bold=True)
    for item in [
        "The current productivity of Copilot (both versions)",
        "Whether Public Gen AI tools provided additional value beyond Copilot",
        "The relative utility of each of the selected Public Gen AI tools",
        "The potential productivity benefits from the Public Gen AI tools",
        "The degree of concern staff have when using the tools",
    ]:
        bullet(doc, item)

    heading(doc, "Overarching findings", 2)
    callout(doc, "Public Gen AI tools delivered clear productivity benefits for many users, but value was not uniform. Copilot already provides meaningful time savings, and public tools offered additional value, particularly for users with lower Copilot access.")
    stat_cards(
        doc,
        "Copilot already saves time, and most public-tool users saw value",
        [
            ("69 min/day", "M365 daily saving", DEWR_DARK_GREY),
            ("34 min/day", "Chat/basic daily saving", DEWR_DARK_GREY),
            ("80%", "Rated tool useful", DEWR_GREEN),
            ("72%", "Continued access", DEWR_DARK_GREEN),
        ],
        "Source: DEWR Public Generative AI Trial survey, 2026. Public-tool results are based on valid public-tool users.",
    )
    for item in [
        "M365 Copilot users saved nearly double the time of Copilot Chat users: 5.7 hours vs 2.8 hours per week.",
        "ChatGPT had the broadest uptake (92%), but Claude showed the strongest usefulness (71%) and continuation intent (63%).",
        "Public tools provided clearer marginal value for Copilot Chat/basic users than for M365 Copilot users.",
        "The most common barrier was lack of integration with internal systems (49%), followed by free-tier prompt limits (41%).",
        "75% of respondents were comfortable using public tools. 11% reported ethical concerns and 3% reported specific security concerns.",
    ]:
        key_finding(doc, item)

    heading(doc, "Recommendations", 2)
    add_paragraph(doc, "The findings reveal several considerations for DEWR in the context of future adoption of Public Generative AI tools.")
    for rec_title, rec_items in [
        ("Targeted access expansion", [
            "Consider providing broader access to public Gen AI tools, prioritising user groups that showed the strongest productivity gains, particularly experienced Gen AI users and Copilot Chat/basic users.",
            "Evaluate whether paid tiers of public tools (especially Claude) would address free-tier limitations that constrained trial benefits.",
        ]),
        ("Capability building", [
            "Offer specialised training reflecting DEWR-specific use cases, building on the finding that prior Gen AI experience was the clearest predictor of stronger outcomes.",
            "Identify and promote Gen AI Champions to share effective practices and encourage adoption.",
        ]),
        ("Integration and workflow", [
            "Investigate integration pathways to reduce the manual effort of transferring public tool outputs into internal systems, which was the most commonly reported barrier.",
            "Conduct workflow analysis across Groups and classifications to identify high-value use cases.",
        ]),
        ("Governance and risk management", [
            "Provide clearer guidance on boundary cases for data classification, where staff reported uncertainty about what was technically allowed but still sensitive.",
            "Proactively monitor the impacts of generative AI on the workforce, including effects on accuracy, over-reliance, and environmental considerations.",
        ]),
    ]:
        heading(doc, rec_title, 3)
        for item in rec_items:
            bullet(doc, item)

    page_break(doc)
    heading(doc, "Evaluation findings", 1, rule=True)
    heading(doc, "1. Copilot productivity baseline", 2)
    callout(doc, "Copilot already delivers material productivity value, but benefits are tiered by access: integrated M365 Copilot produces stronger time savings, deeper engagement and a broader task footprint than Copilot Chat/basic access.")
    value_panel(doc, "M365 Copilot users reported stronger productivity and engagement", [("~2x", "Avg daily time saved"), ("1.8x", "Rated very/extremely useful"), ("1.4x", "Used at least weekly")], note="Note: M365 Copilot n=30; Copilot Chat/basic n=41. Source: DEWR Public Generative AI Trial survey, 2026.")
    heading(doc, "1.1 Integrated Copilot access set a higher productivity baseline", 3)
    add_paragraph(doc, "M365 Copilot users reported average time savings of 69 minutes per day, compared with 34 minutes per day for Copilot Chat/basic users. This means M365 users reported roughly twice the daily time savings of Copilot Chat/basic users.")
    bar_panel(doc, "M365 Copilot users reported about twice the daily time saving", "Access type", [("M365 Copilot - 5.7 hours/week", 69), ("Copilot Chat/basic - 2.8 hours/week", 34)], primary_count=1, note="Note: Weekly equivalents assume a five-day working week. Source: DEWR Public Generative AI Trial survey, 2026.")
    heading(doc, "1.2 Higher time savings were matched by deeper engagement", 3)
    add_paragraph(doc, "The productivity gap was reinforced by engagement signals. M365 Copilot users were more likely to rate Copilot as highly useful, use it weekly, and use it daily or most of the day, suggesting integration is supporting more habitual use.")
    matrix_panel(doc, "M365 Copilot users were more engaged across all usage measures", ["Measure", "M365", "Chat/basic", "M365 lead"], [("Rated very/extremely useful", "67%", "37%", "+30 pts"), ("Used at least weekly", "80%", "56%", "+24 pts"), ("Used daily or most of day", "63%", "27%", "+36 pts")], highlight_col=1, note='Note: M365 Copilot n=30; Copilot Chat/basic n=41. "Useful" means rated very or extremely useful. Source: DEWR Public Generative AI Trial survey, 2026.')
    heading(doc, "1.3 Integrated access supported a broader task footprint", 3)
    add_paragraph(doc, "M365 Copilot users show broader, higher-value task adoption than Copilot Chat users. Both access groups used Copilot for common knowledge-work tasks, but integrated users reported a wider footprint overall, selecting 4.0 task types on average compared with 3.2 for Copilot Chat/basic users. The largest access-type gaps were in research, ideation and meeting preparation.")
    matrix_panel(doc, "M365 Copilot users reported a broader task footprint", ["Task type", "M365", "Chat/basic", "Gap"], [("Summarising", "86%", "76%", "+10 pts"), ("Editing and revision", "72%", "66%", "+6 pts"), ("Drafting", "71%", "61%", "+10 pts"), ("Research, problem solving or ideation", "67%", "44%", "+23 pts"), ("Planning or meeting preparation", "33%", "15%", "+18 pts")], highlight_col=1, note="Note: M365 Copilot n=30; Copilot Chat/basic n=41. Source: DEWR Public Generative AI Trial survey, 2026.")

    page_break(doc)
    heading(doc, "2. Public Gen AI uptake, use and productivity", 2)
    callout(doc, "Public Gen AI tools created clear value for many trial users, but benefits were uneven: strongest for experienced users, staff without integrated Copilot access, and tasks where public tools offered capability or quality beyond Copilot.")
    add_paragraph(doc, "Public Gen AI tools were valued by most trial users, but the benefits were not universal. Overall results were positive, with strong usefulness and continuation signals, while regular use was more moderate.")
    value_panel(doc, "Most public-tool users found value, but regular use was more modest", [("80%", "Rated at least one public tool useful"), ("72%", "Said public tools add value beyond Copilot"), ("72%", "Wanted continued access"), ("53%", "Used public tools at least weekly")], note="Note: Results use valid public-tool user denominators where applicable. Source: DEWR Public Generative AI Trial survey, 2026.")
    heading(doc, "2.1 ChatGPT delivered reach; Claude delivered depth", 3)
    add_paragraph(doc, "ChatGPT was the access point for most trial users, while Claude showed the strongest value signal among users. This suggests the public-tool choice was not simply about uptake; different tools played different roles.")
    matrix_panel(doc, "ChatGPT had the widest reach; Claude had the strongest value signals", ["Measure", "ChatGPT", "Claude", "Gemini"], [("Used during trial", "92%", "67%", "61%"), ("Rated moderately useful or better", "62.5%", "70.7%", "64.9%"), ("Rated very/extremely useful", "28.6%", "46.3%", "29.7%"), ("Wanted continued access", "54%", "63%", "43%")], highlight_col=2, note="Note: Uptake is the share of valid respondents who used each tool. Usefulness and continuation intent are based on users of each respective tool. Source: DEWR Public Generative AI Trial survey, 2026.")
    heading(doc, "2.2 Public tools were mainly used for broad knowledge work", 3)
    add_paragraph(doc, "Use clustered around general knowledge-work tasks rather than specialised or administrative workflows. Research, summarising, editing and drafting were the dominant use cases across the trial.")
    bar_panel(doc, "Public tools were mostly used for broad knowledge-work tasks", "Task type", [("Research, problem solving or ideation", 67), ("Summarising", 66), ("Editing and revision", 62), ("Drafting", 61), ("General administrative tasks", 28), ("Coding or data work", 21), ("Planning and meeting preparation", 20)], primary_count=4, note="Note: Share of valid public-tool users selecting each task type. Source: DEWR Public Generative AI Trial survey, 2026.")
    add_paragraph(doc, "Claude had the strongest research profile (73% of Claude users), while Gemini was relatively stronger for editing and revision. ChatGPT had a more balanced profile across research, summarising, drafting and editing.")
    heading(doc, "2.3 Value was strongest where public tools added marginal capability", 3)
    add_paragraph(doc, "Public tools appeared to provide the highest marginal value where Copilot was less integrated or user capability was higher. Copilot Chat/basic users were more likely to see public tools as adding value over Copilot, despite similar weekly usage.")
    matrix_panel(doc, "Public tools added more marginal value for Copilot Chat/basic users", ["Access type", "Added value beyond Copilot", "Rated useful", "Used weekly+"], [("Copilot Chat/basic", "79%", "82%", "53%"), ("M365 Copilot", "63%", "78%", "52%")], highlight_col=1, note="Source: DEWR Public Generative AI Trial survey, 2026.")
    add_paragraph(doc, "Value also varied by APS/EL level and organisational group. EL users were more likely to report value over Copilot, APS users were more likely to rate at least one public tool as useful, and Workplace Relations recorded the strongest value-over-Copilot result.")
    heading(doc, "2.4 Public tools offered productivity gains beyond Copilot, but value varied", 3)
    add_paragraph(doc, "Claude's value persisted for M365 Copilot users, while ChatGPT added more for Copilot Chat/basic users. Public tools were most positively rated against Copilot for output quality, but the pattern varied by access type.")
    matrix_panel(doc, "ChatGPT added more value for Chat/basic users; Claude stood out for M365 users", ["Access type", "ChatGPT", "Gemini", "Claude"], [("Copilot Chat/basic", "47%", "27%", "39%"), ("M365 Copilot", "27%", "24%", "44%")], highlight_col=1, note='Note: Share rating each public tool better than Copilot, averaged across output quality, time saved, work support, speed and prompt understanding. "Better" includes "a little better" and "much better." Source: DEWR Public Generative AI Trial survey, 2026.')
    heading(doc, "2.5 Integration, limits and reliability constrained scale", 3)
    add_paragraph(doc, "Benefits were constrained less by user interest than by operating conditions. The main limitations were integration with internal systems, free-tier limits and reliability issues.")
    bar_panel(doc, "Integration gaps and free-tier limits were the main barriers to scaling", "Limitation reported", [("Lack of integration with internal systems or Microsoft 365 products", 49), ("Free prompt/request limits", 41), ("Misinterpreted prompts", 34), ("Difficulty with specialised topics", 34), ("Slow responses", 28), ("Fabricated content or hallucinations", 15)], primary_count=2, note="Note: Share of valid public-tool users reporting each limitation. Source: DEWR Public Generative AI Trial survey, 2026.")
    add_paragraph(doc, "Lack of integration was common across all three tools. Free prompt limits were most concentrated in ChatGPT (36%), compared with Claude (20%) and Gemini (5%). Gemini had the highest share reporting difficulty with specialised topics (38%).")

    page_break(doc)
    heading(doc, "3. Concerns, risks and safeguards", 2)
    callout(doc, "Most valid users were comfortable and reported security concerns were rare. Survey results suggest safety communications and splash screens supported cautious use, while data-handling risks were mainly visible through copy/paste behaviour and user judgement.")
    heading(doc, "3.1 Comfort was high, but risk signals were not absent", 3)
    add_paragraph(doc, "Staff comfort was relatively high, but the concern profile was not negligible. Three-quarters of respondents were comfortable using public tools, while one-quarter remained uncomfortable and smaller shares reported ethical or security concerns.")
    value_panel(doc, "Most respondents were comfortable, but one in four were not", [("75%", "Comfortable or very comfortable using public tools"), ("25%", "Uncomfortable using public tools"), ("11%", "Ethical concerns encountered"), ("3%", "Reported specific security concerns")], note="Source: DEWR Public Generative AI Trial survey, 2026.")
    heading(doc, "3.2 Concerns centred on practical operating boundaries", 3)
    add_paragraph(doc, "Open-text concerns were broader than reported incidents and centred on the practical conditions for safe use. The most common themes were tool access and integration, data handling, and output accuracy.")
    concern_panel(doc)

    page_break(doc)
    heading(doc, "3.3 Safety communications were rated effective by most valid users", 3)
    add_paragraph(doc, "The survey provides the clearest positive evidence for the communications layer. Around seven in ten valid users rated the introductory email and splash screens as moderately or highly effective. Reported security concerns were rare, and no valid user rated upload blockers as ineffective.")
    value_panel(doc, "Most valid users rated email and splash screens effective", [("74%", "Introductory email moderately or highly effective"), ("72%", "Splash screens moderately or highly effective"), ("67%", "Rated both email and splash screens effective"), ("0%", "Rated upload blockers ineffective")], primary_count=2, note="Note: Percentages use the valid public-tool user denominator where available (n=60-61 depending on skipped items). Upload-blocker results reflect ratings and visibility, not trigger frequency. Source: DEWR Public Generative AI Trial survey, 2026.")
    heading(doc, "3.4 Positive safety communications aligned with higher comfort", 3)
    add_paragraph(doc, "Users who rated both the introductory email and splash screens positively were more likely to feel comfortable using public tools. This does not prove causation, but it supports the interpretation that communications helped users understand trial boundaries and operate with more confidence.")
    bar_panel(doc, "Comfort was higher among users who found both safety communications effective", "User group", [("Rated both email and splash screens effective", 82.5), ("Did not rate both channels effective", 61.9)], primary_count=1, note="Note: Comfort means comfortable or very comfortable using public tools. Source: DEWR Public Generative AI Trial survey, 2026.")
    page_break(doc)
    heading(doc, "3.5 Data handling relied more on user judgement than visible upload blocking", 3)
    add_paragraph(doc, "Most valid users did not notice upload blockers, while copying and pasting information was more common than document upload. This suggests the survey captured stronger evidence of user behaviour and judgement than of technical blocker performance. No valid user rated upload blockers ineffective, but the low visibility of blockers means the survey provides limited direct evidence about how often they were triggered.")
    bar_panel(doc, "Copy/paste was more common than document upload; blockers had low visibility", "Measure", [("Copied and pasted information", 70.5), ("Uploaded documents", 42.6), ("Noticed/rated upload blockers effective", 16.7), ("Rated upload blockers ineffective", 0)], primary_count=1, note="Note: 43 of 61 valid users copied/pasted information; 26 of 61 uploaded documents; 10 of 60 rated upload blockers effective; no valid user rated blockers ineffective. Source: DEWR Public Generative AI Trial survey, 2026.")
    heading(doc, "3.6 Staff were making judgement calls at the edge of guidance", 3)
    add_paragraph(doc, "The strongest qualitative signal was not a single incident category, but uncertainty at the boundary of policy and practice. Some users described having to decide in the moment whether information was appropriate for public tools, how to validate outputs, and when Copilot's APS-specific safeguards mattered.")
    judgement_panel(doc)
    add_quote(doc, "There were a few occasions where I had to question myself before hitting send on a prompt, project management artifacts could be hard to classify and will require the user to make a judgement call in the moment.")
    add_quote(doc, "I chose not to upload some documents & use the AI because even though they were within the allowed classifications, they contained information that could be deemed commercially confidential.")
    add_quote(doc, "When asking the AI tools about the State reservations on CEDAW, Copilot gave an answer in very neutral language. The other three AI tools all gave what I would describe as far less sensitive responses... This highlighted the value of the extra APS sensitivities that have been put into Copilot chat.")
    add_paragraph(doc, "Environmental sustainability, over-reliance and broader workforce effects were raised less frequently than data and access issues, but they were still part of the broader concern profile reported by users.")
    add_quote(doc, "No indication of how resource intensive AI LLM are / the environmental impacts of using these tools. I am not confident that the risks to climate change or the environment have been fully considered by government.")

    page_break(doc)
    heading(doc, "Approach and methodology", 1, rule=True)
    heading(doc, "Goal and scope", 3)
    add_paragraph(doc, "The trial assessed whether providing DEWR staff with access to publicly available Generative AI tools would deliver productivity benefits beyond those already provided by Microsoft Copilot. The trial ran for a defined period in January 2026 with approximately 200 randomly selected and stratified employees.")
    heading(doc, "Sampling", 3)
    for item in ["5% of DEWR employees were selected for the trial.", "Selection was randomised and stratified across Groups, Copilot access types, and APS levels.", "All trial participants retained their existing Copilot access throughout."]:
        bullet(doc, item)
    heading(doc, "Survey design and response", 3)
    for item in ["A voluntary post-trial survey was administered to all trial participants.", "104 staff completed the survey (approximately 52% response rate).", "Respondents: 52 APS staff (50%) and 52 Executive Level staff (50%).", "Respondents came from all five departmental groups.", "71 respondents (68.3%) used at least one public Gen AI tool during the trial.", "30 respondents reported having M365 Copilot access."]:
        bullet(doc, item)
    matrix_panel(doc, "Most respondents had at least some prior Gen AI experience", ["Experience level", "Count", "Share"], [("No prior experience", "7", "6.7%"), ("Basic familiarity", "33", "31.7%"), ("Some experience", "35", "33.7%"), ("Experienced", "24", "23.1%"), ("Highly experienced", "5", "4.8%")], note="Note: n=104; shares may not total 100 due to rounding. Source: DEWR Public Generative AI Trial survey, 2026.")

    page_break(doc)
    heading(doc, "Appendix", 1, rule=True)
    heading(doc, "Survey questionnaire", 3)
    add_paragraph(doc, "The survey comprised mandatory and optional questions across six sections. Questions requiring an answer are denoted with an asterisk (*). Skip logic directed participants to later questions based on their responses.")
    for section, items in [
        ("Section 1: General questions (Q1-Q11)", ["APS level, Group, Job Family, Job Title", "Prior Gen AI experience level", "Whether the participant used public Gen AI tools during the trial", "Reasons for non-use (if applicable)", "Frequency of public tool use, experience improvement, document upload and copy-paste behaviour"]),
        ("Section 2: Copilot (Q12-Q17)", ["M365 Copilot access status", "Copilot usage frequency, task types, overall usefulness", "Time saved by activity area and across the average workday"]),
        ("Section 3: ChatGPT (Q18-Q25)", ["Whether ChatGPT was used, frequency, task types, usefulness", "Comparison with Copilot across five dimensions", "Additional value beyond Copilot by task type", "Limitations experienced and continuation intent"]),
        ("Section 4: Gemini (Q26-Q33)", ["Same structure as ChatGPT section"]),
        ("Section 5: Claude (Q34-Q41)", ["Same structure as ChatGPT section"]),
        ("Section 6: Final questions (Q42-Q48)", ["Overall comfort using public Gen AI tools", "Value of public tools over and above Copilot", "Ethical and security concerns encountered", "Open-text feedback on concerns and additional comments"]),
    ]:
        heading(doc, section, 3)
        for item in items:
            bullet(doc, item)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_doc()
    print(f"Report generated: {path}")
